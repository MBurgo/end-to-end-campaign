import streamlit as st
import json
import time
import re
import pathlib
from io import BytesIO
from textwrap import dedent
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from openai import OpenAI

# ────────────────────────────────────────────────────────────
# 1. PAGE CONFIG & CSS
# ────────────────────────────────────────────────────────────
st.set_page_config(page_title="Foolish Campaign Studio", page_icon="🃏", layout="wide")

st.markdown("""
<style>
    .stButton>button { width: 100%; border-radius: 8px; height: 3em; }
    .chat-bubble { padding: 15px; border-radius: 10px; margin-bottom: 10px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
    .believer-bubble { background-color: #e3f6d8; border-left: 5px solid #43B02A; color: #1f1f1f; }
    .skeptic-bubble { background-color: #fce8e6; border-left: 5px solid #d93025; color: #1f1f1f; }
    h1, h2, h3 { color: #1f1f1f; }
    .instruction-box { background-color: #f0f2f6; padding: 15px; border-radius: 8px; border-left: 5px solid #5865F2; margin-bottom: 20px; }
    .step-header { font-weight: bold; font-size: 1.1em; margin-bottom: 0.5em; color: #5865F2; }
</style>
""", unsafe_allow_html=True)

# ────────────────────────────────────────────────────────────
# 2. CONFIGURATION & DATA LOADING
# ────────────────────────────────────────────────────────────

@st.cache_data
def load_configs():
    # Load Traits
    try:
        traits = json.loads(pathlib.Path("traits_config.json").read_text())
    except Exception:
        st.error("🚨 Missing 'traits_config.json'.")
        traits = {}

    # Load Personas
    try:
        raw_p = json.loads(pathlib.Path("personas.json").read_text())
        flat_personas = []
        for group in raw_p.get("personas", []):
            segment = group.get("segment", "Unknown")
            for gender in ["male", "female"]:
                if gender in group:
                    p = group[gender]
                    p["id"] = f"{p['name']} ({segment})"
                    p["segment"] = segment
                    flat_personas.append(p)
        return traits, flat_personas
    except Exception as e:
        st.error(f"🚨 Error loading 'personas.json': {e}")
        return traits, []

TRAIT_CFG, ALL_PERSONAS = load_configs()

# ────────────────────────────────────────────────────────────
# 3. AI ENGINE HANDLER (Hybrid with Self-Healing)
# ────────────────────────────────────────────────────────────

def query_openai(messages, temperature=0.7, json_mode=False):
    """Specialized handler for Personas (GPT-4o)."""
    if "openai_api_key" not in st.secrets:
        st.error("⚠️ OpenAI API Key missing in .streamlit/secrets.toml")
        return None
        
    client = OpenAI(api_key=st.secrets.openai_api_key)
    kwargs = {"model": "gpt-4o", "temperature": temperature}
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    
    # Simple retry loop for OpenAI
    for attempt in range(3):
        try:
            resp = client.chat.completions.create(messages=messages, **kwargs)
            return resp.choices[0].message.content
        except Exception as e:
            if attempt == 2: # Last attempt
                st.error(f"OpenAI Error: {e}")
                return None
            time.sleep(1) # Wait 1s before retry

def query_gemini(messages, temperature=0.7, json_mode=False):
    """Specialized handler for Copy & Strategy with Robust Fallback."""
    if "google_api_key" not in st.secrets:
        st.error("⚠️ Google API Key missing in .streamlit/secrets.toml")
        return None
        
    genai.configure(api_key=st.secrets.google_api_key)
    
    # FIX: Default to None if no system message is found
    sys_msg = next((m['content'] for m in messages if m['role'] == 'system'), None)
    
    history = [m['content'] for m in messages if m['role'] != 'system']
    prompt = "\n\n".join(history)

    safety_config = {
        HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
    }

    config = genai.GenerationConfig(
        temperature=temperature,
        response_mime_type="application/json" if json_mode else "text/plain"
    )
    
    # ---------------------------------------------------------
    # SELF-HEALING MODEL SELECTOR
    # ---------------------------------------------------------
    known_models = [
        "gemini-2.5-flash",          # BEST OPTION
        "gemini-2.5-flash-lite",     # BACKUP
        "gemini-2.0-flash",          # LEGACY
        "gemini-3-flash-preview"     # EXPERIMENTAL
    ]

    last_error = None

    for model_name in known_models:
        for attempt in range(3):
            try:
                # Only pass system_instruction if it exists (not None)
                if sys_msg:
                    model = genai.GenerativeModel(model_name=model_name, system_instruction=sys_msg)
                else:
                    model = genai.GenerativeModel(model_name=model_name)
                    
                resp = model.generate_content(prompt, generation_config=config, safety_settings=safety_config)
                return resp.text
                
            except Exception as e:
                error_str = str(e)
                last_error = error_str
                
                if "429" in error_str:
                    time.sleep(2 ** (attempt + 1)) # Exponential backoff
                    continue
                if "404" in error_str or "not found" in error_str.lower():
                    break 
                if attempt == 2:
                    break
        else:
            continue
        break

    if last_error:
        st.error(f"🚨 All Gemini models failed. Last error: {last_error}")
    return None

# ────────────────────────────────────────────────────────────
# 4. HELPER FUNCTIONS
# ────────────────────────────────────────────────────────────

def trait_rules_builder(scores):
    rules = []
    for name, score in scores.items():
        cfg = TRAIT_CFG.get(name)
        if not cfg: continue
        if score >= cfg["high_threshold"]: rules.append(cfg["high_rule"])
        elif score <= cfg["low_threshold"]: rules.append(cfg["low_rule"])
        elif "mid_rule" in cfg: rules.append(cfg["mid_rule"])
    return rules

def create_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for line in text.split('\n'):
        if line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=2)
        elif line.strip(): doc.add_paragraph(line.replace('**', '').replace('* ', ''))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ────────────────────────────────────────────────────────────
# 5. SESSION STATE
# ────────────────────────────────────────────────────────────
defaults = {
    "current_draft": "",
    "internal_plan": "",
    "debate_transcript": [],
    "critique_summary": "",
    "optimized_copy": "",
}
for k, v in defaults.items():
    if k not in st.session_state: st.session_state[k] = v

# ────────────────────────────────────────────────────────────
# 6. SIDEBAR
# ────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("🎛️ Campaign Settings")
    st.info("🤖 **Hybrid Intelligence Active**\n\n• **Writer:** Gemini 2.5 Flash\n• **Personas:** OpenAI GPT-4o\n• **Strategist:** Gemini 2.5 Flash")
    st.markdown("---")
    
    st.subheader("🎯 Primary Audience")
    persona_names = [p["id"] for p in ALL_PERSONAS]
    target_persona_name = st.selectbox("Select Avatar", persona_names, index=0 if persona_names else None)
    target_persona = next((p for p in ALL_PERSONAS if p["id"] == target_persona_name), None)

    st.markdown("---")

    with st.expander("🎚️ Voice & Tone Control", expanded=True):
        trait_scores = {
            "Urgency": st.slider("Urgency (Scarcity)", 1, 10, 7),
            "Data_Richness": st.slider("Specificity (4 U's)", 1, 10, 6),
            "Social_Proof": st.slider("Social Proof", 1, 10, 6),
            "Comparative_Framing": st.slider("Contrast Principle", 1, 10, 5),
            "Conversational_Tone": st.slider("Liking/Relatability", 1, 10, 8),
            "FOMO": st.slider("Fear (Missing Out)", 1, 10, 5),
        }

# ────────────────────────────────────────────────────────────
# 7. MAIN INTERFACE
# ────────────────────────────────────────────────────────────
st.title("🃏 Foolish Campaign Studio")

tab_write, tab_test, tab_refine = st.tabs(["1️⃣ Write Draft", "2️⃣ Stress Test", "3️⃣ Optimize"])

# ============================================================
# TAB 1: THE DRAFTER
# ============================================================
with tab_write:
    st.markdown("""
    <div class="instruction-box">
        <div class="step-header">Step 1: The Foolish Brief</div>
        Fill in the details. The AI uses the <strong>"Raising Response"</strong> curriculum to write 
        emotionally driven direct-response copy.
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📋 Campaign Brief")
        hook = st.text_area("The 'Big Idea' (Hook)", "e.g. 3 AI Stocks better than Nvidia. Urgent Buy Alert!", height=100, help="Must be Timely, Useful, and Unique.")
        details = st.text_area("Product / Offer Details", "e.g. Subscription to Share Advisor. $199/year.", height=150)
        
        c1, c2 = st.columns(2)
        length_mode = c1.selectbox("Length", ["Short (200w)", "Medium (500w)", "Long (1000w)"])
        copy_type = c2.selectbox("Format", ["📧 Email", "📝 Sales Page"])
        
        # New: Select the Lede Type from the PDF 
        lede_type = st.selectbox("Lede Strategy", [
            "The Interrupting Idea (Startle the reader)",
            "The Shocker (Go against the grain)",
            "The News (Timely/Relevant)",
            "The Preview (Tease the benefit)",
            "The Story (Fact-packed narrative)",
            "The Quotation (Hyper-relevant)"
        ])
        
        if st.button("✨ Generate Foolish Copy", type="primary"):
            if not hook or not details:
                st.warning("Please enter a Hook and Details.")
            else:
                rules = trait_rules_builder(trait_scores)
                target_desc = f"{target_persona['name']}, {target_persona['age']}yo {target_persona['occupation']} ({target_persona['segment']})" if target_persona else "General Investor"
                
                # ENHANCED SYSTEM PROMPT: Injecting "Raising Response" DNA
                sys_prompt = dedent(f"""
                You are a Master Direct Response Copywriter for The Motley Fool. 
                You follow the principles of the "Raising Response" course.

                ### CORE PHILOSOPHY
                1. **AIDA Framework:** You must grab Attention, generate Interest, stimulate Desire, and force Action.
                2. **The 4 U's:** Your headline must be Urgent, Unique, Useful, and Ultra-Specific.
                3. **Emotional Drivers:** You must target one of the 7 Drivers: Lust, Escape, Esteem, Fear, Guilt, Affinity, or Greed[cite: 374].
                
                ### STYLE RULES [cite: 1197]
                - Use small words. Avoid "poetic" or "clever" copy.
                - Write in the second person ("You").
                - Use active verbs and the present tense.
                - Be conversational but authoritative.
                
                ### TARGET AUDIENCE
                {target_desc}
                
                ### TONE & BEHAVIOR RULES
                {chr(10).join(rules)}
                
                ### STRUCTURE
                - Use Markdown headings.
                - End with compliance line: *Past performance is not a reliable indicator of future results.*
                """)
                
                user_prompt = dedent(f"""
                Write a {copy_type} ({length_mode}).
                
                **THE BIG IDEA:** {hook}
                **LEDE STRATEGY:** {lede_type}
                **OFFER DETAILS:** {details}
                
                **TASK:**
                1. Create a bulleted "Internal Plan" identifying the specific Emotional Driver and how you will apply the 4 U's.
                2. Write the copy.
                
                OUTPUT JSON: {{ "plan": "...", "copy": "..." }}
                """)
                
                with st.spinner("Gemini is crafting your copy..."):
                    msgs = [{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_prompt}]
                    res = query_gemini(msgs, json_mode=True)
                    
                    if res:
                        try:
                            json_match = re.search(r"\{[\s\S]*\}", res)
                            if json_match:
                                clean_json = json_match.group(0)
                                data = json.loads(clean_json)
                                st.session_state.current_draft = data["copy"]
                                st.session_state.internal_plan = data["plan"]
                                st.session_state.debate_transcript = [] 
                                st.session_state.critique_summary = ""
                                st.rerun()
                            else:
                                st.error("AI returned text but no JSON format found.")
                                st.write(res)
                        except Exception as e:
                            st.error(f"⚠️ Parsing Error: {e}")
                            st.code(res)

    with col2:
        st.subheader("📝 Draft Preview")
        if st.session_state.current_draft:
            with st.expander("View AI's Strategy (The Plan)", expanded=True):
                st.info(st.session_state.internal_plan)
            st.markdown(st.session_state.current_draft)
        else:
            st.info("👈 Fill out the brief and click 'Generate Foolish Copy' to start.")

# ============================================================
# TAB 2: THE SIMULATOR
# ============================================================
with tab_test:
    st.markdown("""
    <div class="instruction-box">
        <div class="step-header">Step 2: The Focus Group</div>
        Select two diverse personas to debate your draft. 
        <strong>The Believer</strong> focuses on Opportunity (Greed/Lust). 
        <strong>The Skeptic</strong> focuses on Risk (Fear/Guilt).
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.current_draft:
        st.warning("⚠️ Please generate a draft in Tab 1 first.")
    else:
        c_a, c_b, c_c = st.columns([1, 1, 1])
        with c_a:
            idx_1 = persona_names.index(target_persona_name) if target_persona_name in persona_names else 0
            p1_name = st.selectbox("Role: The Believer", persona_names, index=idx_1, key="sim_p1")
        with c_b:
            p2_name = st.selectbox("Role: The Skeptic", persona_names, index=(idx_1 + 1) % len(persona_names), key="sim_p2")
        with c_c:
            st.write("") 
            st.write("") 
            start_sim = st.button("▶️ Start Simulation (GPT-4o)")
            
        if start_sim:
            p1 = next(p for p in ALL_PERSONAS if p["id"] == p1_name)
            p2 = next(p for p in ALL_PERSONAS if p["id"] == p2_name)
            draft_text = st.session_state.current_draft
            transcript = []
            
            prompt_base = f"You are participating in a marketing focus group. React to the copy below. Be conversational."
            
            # 1. The Believer
            role_1 = f"You are {p1['name']}, {p1['age']}. Bio: {p1['narrative']}. Values: {p1['values']}. You are OPTIMISTIC. You focus on the 'Upside' and the 'Dream'."
            msg_1_prompt = f"{prompt_base}\nTEXT TO REVIEW:\n'{draft_text}'\n\nWhat excites you? Does the 'Big Idea' hook you?"
            
            with st.status("Running Simulation...") as status:
                status.write(f"🎤 {p1['name']} is reading...")
                res_1 = query_openai([{"role":"system", "content": role_1}, {"role":"user", "content": msg_1_prompt}])
                transcript.append({"name": p1["name"], "role": "Believer", "text": res_1})
                
                # 2. The Skeptic
                status.write(f"🤔 {p2['name']} is critiquing...")
                role_2 = f"You are {p2['name']}, {p2['age']}. Bio: {p2['narrative']}. Values: {p2['values']}. You are SKEPTICAL. You focus on 'Risk' and 'Objections'."
                msg_2_prompt = f"{prompt_base}\nTEXT TO REVIEW:\n'{draft_text}'\n\n{p1['name']} just said: '{res_1}'.\nTell them why they are wrong. Find the holes in the argument."
                
                res_2 = query_openai([{"role":"system", "content": role_2}, {"role":"user", "content": msg_2_prompt}])
                transcript.append({"name": p2["name"], "role": "Skeptic", "text": res_2})
                
                # 3. Rebuttal
                status.write(f"🗣️ {p1['name']} is responding...")
                msg_3_prompt = f"You just heard {p2['name']} say: '{res_2}'. Do you agree with their worry, or do you still want to buy?"
                
                res_3 = query_openai([{"role":"system", "content": role_1}, {"role":"user", "content": msg_3_prompt}])
                transcript.append({"name": p1["name"], "role": "Believer", "text": res_3})
                
                st.session_state.debate_transcript = transcript
                status.update(label="Focus Group Complete!", state="complete", expanded=False)

        if st.session_state.debate_transcript:
            st.divider()
            for turn in st.session_state.debate_transcript:
                style_class = "believer-bubble" if turn["role"] == "Believer" else "skeptic-bubble"
                st.markdown(f"""
                <div class="chat-bubble {style_class}">
                    <strong>{turn['name']} ({turn['role']}):</strong><br>{turn['text']}
                </div>
                """, unsafe_allow_html=True)

# ============================================================
# TAB 3: THE OPTIMIZER
# ============================================================
with tab_refine:
    st.markdown("""
    <div class="instruction-box">
        <div class="step-header">Step 3: Strategic Polish</div>
        The AI Strategist will now analyze the debate to find the <strong>"Trust Gap"</strong> 
        and rewrite the copy to better handle objections.
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.debate_transcript:
        st.info("⚠️ Please run the Focus Group in Tab 2 to unlock optimization.")
    else:
        col_L, col_R = st.columns([1, 1])
        
        with col_L:
            st.subheader("📊 Insight Extraction")
            if st.button("🧠 Analyze Transcript (Gemini)"):
                transcript_text = "\n".join([f"{t['role']}: {t['text']}" for t in st.session_state.debate_transcript])
                
                analysis_prompt = dedent(f"""
                You are a Senior Marketing Strategist using the 'Raising Response' framework.
                Analyze this debate transcript regarding a piece of copy.
                TRANSCRIPT:
                {transcript_text}
                
                Identify:
                1. The main "Trust Gap" (what made the Skeptic doubt)[cite: 89].
                2. The main "Hook Strength" (what the Believer liked).
                3. Three specific actionable edits to improve the copy using 'Objection Removal' techniques[cite: 1387].
                """)
                
                with st.spinner("Analyzing psychology..."):
                    st.session_state.critique_summary = query_gemini([{"role":"user", "content": analysis_prompt}])
            
            if st.session_state.critique_summary:
                st.success(st.session_state.critique_summary)

        with col_R:
            st.subheader("✨ Final Polish")
            if st.session_state.critique_summary:
                if st.button("✍️ Rewrite Copy (Gemini)"):
                    rewrite_prompt = dedent(f"""
                    You are an Expert Editor.
                    ORIGINAL COPY:
                    {st.session_state.current_draft}
                    
                    CRITIQUE TO ADDRESS:
                    {st.session_state.critique_summary}
                    
                    TASK:
                    Rewrite the copy to address the critique while maintaining the original high-energy tone.
                    Strengthen the proof points where the Skeptic was doubtful.
                    Ensure the headline adheres to the 4 U's (Urgent, Unique, Useful, Ultra-Specific).
                    Output ONLY the final copy.
                    """)
                    
                    with st.spinner("Polishing final draft..."):
                        st.session_state.optimized_copy = query_gemini([{"role":"user", "content": rewrite_prompt}])
            
            if st.session_state.optimized_copy:
                st.markdown(st.session_state.optimized_copy)
                docx = create_docx(st.session_state.optimized_copy)
                st.download_button("📥 Download Final DOCX", docx, "final_campaign.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
